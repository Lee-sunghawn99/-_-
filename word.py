from docx import Document
import argparse
from io import BytesIO

parser = argparse.ArgumentParser()

parser.add_argument(
    "--prompt", type=str, nargs="?", default="happy puppy is running with friends"
)

parser.add_argument(
    "--seed", type=int, nargs="?", default=10000
)


user_input = parser.parse_args()
user_input.seed2 = user_input.seed + 1

title = f'''my dream-{user_input.prompt}'''
image1 = f'''seed_{user_input.seed}_00000'''
image2 = f'''seed_{user_input.seed2}_00001'''
text_with_underscores = user_input.prompt.replace(" ", "_")

# create document object
document = Document()

image_path1 = f'''/home/lee99/stable-diffusion/outputs/txt2img-samples/{text_with_underscores}/{image1}.png'''
with open(image_path1, "rb") as f:
    image_data1 = f.read()
image_path2 = f'''/home/lee99/stable-diffusion/outputs/txt2img-samples/{text_with_underscores}/{image2}.png'''
with open(image_path2, "rb") as f:
    image_data2 = f.read()

text_path = f'''/home/lee99/문서/AI_bigdata/{text_with_underscores}'''
with open(text_path, "r") as f:
    text_data = f.read()

image_data1 = BytesIO(image_data1)
image_data2 = BytesIO(image_data2)

document.add_heading(f'''my dream-{user_input.prompt}''')
document.add_picture(image_data1)
document.add_paragraph(text_data)
document.add_picture(image_data2)

document.save(f'''/home/lee99/문서/AI_bigdata/dream_story/my_dream-{text_with_underscores}.docx''')
